# =============================================================================
# IMPORTS
# =============================================================================
import sys

st.write(sys.version)

import streamlit as st
import pandas as pd
import numpy as np

import plotly.graph_objects as go
try:
    import plotly
    print("Plotly carregado:", plotly.__version__)

    import plotly.express as px

except Exception as e:
    st.error(f"Erro Plotly: {e}")
    raise
import requests
import zipfile
from io import BytesIO
from scipy.stats import zscore
from docx import Document
from openpyxl import Workbook

# =============================================================================
# CONFIGURAÇÃO DA PÁGINA
# =============================================================================

st.set_page_config(
    page_title="AgroClimate AI",
    page_icon="🌱",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =============================================================================
# CSS GLOBAL
# =============================================================================

st.markdown(
    """
    <style>

    .main {
        padding-top: 1rem;
    }

    .stMetric {
        border-radius: 12px;
        padding: 10px;
    }

    .block-container {
        padding-top: 1rem;
    }

    footer {
        visibility: hidden;
    }

    </style>
    """,
    unsafe_allow_html=True
)

# =============================================================================
# SESSION STATE
# =============================================================================

if "usuario" not in st.session_state:
    st.session_state["usuario"] = {}

if "dados_salvos" not in st.session_state:
    st.session_state["dados_salvos"] = False

# =============================================================================
# TÍTULO PRINCIPAL
# =============================================================================

st.title("🌱 AgroClimate AI")
st.caption(
    "Sistema Inteligente para Processamento, Análise e Interpretação de Dados Climáticos e Agrícolas"
)

st.markdown("---")

# =============================================================================
# CRIAÇÃO DAS ABAS
# =============================================================================

tab0, tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9 = st.tabs(
    [
        "🏠 Início",
        "📂 Importação",
        "🧹 Tratamento",
        "📅 Consolidação",
        "📈 Estatística",
        "📊 Gráficos",
        "🤖 IA",
        "📥 Exportação",
        "🌱 Indicadores"
        "🔍 Qualidade"
    ]
)

# =============================================================================
# ABA 0 - INÍCIO
# =============================================================================

with tab0:

    st.subheader("🏠 Cadastro do Projeto")

    st.markdown(
        """
        Preencha os dados abaixo para identificação do projeto,
        pesquisador e instituição.
        """
    )

    col1, col2 = st.columns(2)

    with col1:

        nome_projeto = st.text_input(
            "Nome do Projeto"
        )

        pesquisador = st.text_input(
            "Pesquisador Responsável"
        )

        instituicao = st.text_input(
            "Instituição"
        )

    with col2:

        email = st.text_input(
            "E-mail"
        )

        cidade = st.text_input(
            "Cidade"
        )

        estado = st.text_input(
            "Estado"
        )

    observacoes = st.text_area(
        "Observações"
    )

    if st.button(
        "💾 Salvar Cadastro",
        use_container_width=True
    ):

        st.session_state["usuario"] = {

            "Projeto": nome_projeto,
            "Pesquisador": pesquisador,
            "Instituição": instituicao,
            "Email": email,
            "Cidade": cidade,
            "Estado": estado,
            "Observações": observacoes

        }

        st.session_state["dados_salvos"] = True

        st.success(
            "Cadastro salvo com sucesso."
        )

    st.markdown("---")

    if st.session_state["dados_salvos"]:

        st.subheader("📋 Informações Salvas")

        st.json(
            st.session_state["usuario"]
        )

    st.markdown("---")

    st.info(
        """
        Fluxo recomendado:

        1️⃣ Importar dados

        2️⃣ Tratar dados

        3️⃣ Consolidar dados

        4️⃣ Gerar estatísticas

        5️⃣ Criar gráficos

        6️⃣ Gerar relatório com IA

        7️⃣ Exportar resultados

        8️⃣ Avaliar indicadores agrometeorológicos
        """
    )
# =============================================================================
# FUNÇÕES DE IMPORTAÇÃO E DIAGNÓSTICO
# =============================================================================

def detectar_delimitador(arquivo):
    """
    Detecta automaticamente o delimitador do arquivo.
    """

    try:

        arquivo.seek(0)

        amostra = arquivo.read(5000)

        if isinstance(amostra, bytes):
            amostra = amostra.decode(
                "utf-8",
                errors="ignore"
            )

        delimitadores = [
            ",",
            ";",
            "\t",
            "|"
        ]

        contagem = {}

        for d in delimitadores:
            contagem[d] = amostra.count(d)

        return max(
            contagem,
            key=contagem.get
        )

    except:

        return ";"


# =============================================================================
# DETECÇÃO DE ENCODING
# =============================================================================

def detectar_encoding(arquivo):

    encodings = [
        "utf-8",
        "latin1",
        "ISO-8859-1",
        "cp1252"
    ]

    for enc in encodings:

        try:

            arquivo.seek(0)

            conteudo = arquivo.read(3000)

            if isinstance(conteudo, bytes):
                conteudo.decode(enc)

            return enc

        except:
            continue

    return "latin1"


# =============================================================================
# REMOVER COLUNAS DUPLICADAS
# =============================================================================

def remover_colunas_duplicadas(df):

    df = df.loc[
        :,
        ~df.columns.duplicated()
    ]

    return df


# =============================================================================
# CONVERTER COLUNAS NUMÉRICAS
# =============================================================================

def converter_colunas_numericas(df):

    df2 = df.copy()

    for col in df2.columns:

        try:

            df2[col] = (
                df2[col]
                .astype(str)
                .str.replace(",", ".", regex=False)
            )

            df2[col] = pd.to_numeric(
                df2[col],
                errors="ignore"
            )

        except:
            pass

    return df2


# =============================================================================
# DETECTAR COLUNAS DE DATA
# =============================================================================

def detectar_colunas_data(df):

    candidatas = []

    palavras = [
        "data",
        "date",
        "datetime",
        "tempo",
        "hora",
        "timestamp"
    ]

    for col in df.columns:

        nome = str(col).lower()

        if any(
            p in nome
            for p in palavras
        ):

            candidatas.append(col)

    return candidatas


# =============================================================================
# CONVERTER DATAS
# =============================================================================

def converter_datas(df):

    df2 = df.copy()

    datas = detectar_colunas_data(df2)

    for col in datas:

        try:

            df2[col] = pd.to_datetime(
                df2[col],
                errors="coerce",
                dayfirst=True
            )

        except:
            pass

    return df2


# =============================================================================
# IDENTIFICAÇÃO DO TIPO DA PLANILHA
# =============================================================================

def identificar_tipo_planilha(df):

    nomes = " ".join(
        df.columns.astype(str)
    ).lower()

    termos_climaticos = [

        "temp",
        "temperatura",
        "umidade",
        "vento",
        "rad",
        "radiação",
        "radiacao",
        "chuva",
        "precip",
        "eto",
        "evap"

    ]

    score = sum(
        termo in nomes
        for termo in termos_climaticos
    )

    if score >= 2:
        return "Climática"

    return "Genérica"


# =============================================================================
# LEITOR UNIVERSAL
# =============================================================================

def ler_planilha_universal(arquivo):

    nome = arquivo.name.lower()

    # Excel
    if nome.endswith(
        (".xlsx", ".xls")
    ):

        df = pd.read_excel(
            arquivo
        )

    # Texto
    else:

        delimitador = detectar_delimitador(
            arquivo
        )

        encoding = detectar_encoding(
            arquivo
        )

        arquivo.seek(0)

        df = pd.read_csv(
            arquivo,
            sep=delimitador,
            encoding=encoding,
            on_bad_lines="skip"
        )

    df = remover_colunas_duplicadas(df)

    df = converter_colunas_numericas(df)

    df = converter_datas(df)

    return df


# =============================================================================
# ABA 1 - IMPORTAÇÃO E DIAGNÓSTICO
# =============================================================================

with tab1:

    st.subheader(
        "📂 Importação Inteligente"
    )

    st.success(
        """
        Compatível com:

        CSV • TXT • DAT • XLS • XLSX
        """
    )

    uploaded_file = st.file_uploader(

        "Selecione a planilha",

        type=[
            "csv",
            "txt",
            "dat",
            "xls",
            "xlsx"
        ]
    )

    if uploaded_file is not None:

        try:

            df = ler_planilha_universal(
                uploaded_file
            )

            st.session_state[
                "df_original"
            ] = df

            tipo_planilha = (
                identificar_tipo_planilha(df)
            )

            st.markdown("---")

            col1, col2, col3, col4 = (
                st.columns(4)
            )

            with col1:

                st.metric(
                    "Linhas",
                    f"{len(df):,}"
                )

            with col2:

                st.metric(
                    "Colunas",
                    len(df.columns)
                )

            with col3:

                st.metric(
                    "Valores Ausentes",
                    int(
                        df.isna()
                        .sum()
                        .sum()
                    )
                )

            with col4:

                st.metric(
                    "Tipo",
                    tipo_planilha
                )

            st.markdown("---")

            st.subheader(
                "👁 Pré-visualização"
            )

            st.dataframe(
                df.head(100),
                use_container_width=True
            )

            st.markdown("---")

            st.subheader(
                "📋 Estrutura das Colunas"
            )

            tipos = pd.DataFrame({

                "Coluna":
                    df.columns,

                "Tipo":
                    df.dtypes.astype(str)

            })

            st.dataframe(
                tipos,
                use_container_width=True
            )

            st.markdown("---")

            st.subheader(
                "⚠ Valores Ausentes"
            )

            faltantes = pd.DataFrame({

                "Coluna":
                    df.columns,

                "Faltantes":
                    df.isna().sum(),

                "Percentual (%)":
                    (
                        df.isna().sum()
                        / len(df)
                        * 100
                    ).round(2)

            })

            st.dataframe(
                faltantes,
                use_container_width=True
            )

            st.success(
                "Arquivo carregado com sucesso."
            )

        except Exception as erro:

            st.error(
                f"Erro ao carregar arquivo: {erro}"
            )
# =============================================================================
# FUNÇÕES DE TRATAMENTO DE DADOS
# =============================================================================

def obter_colunas_numericas(df):

    return df.select_dtypes(
        include=np.number
    ).columns.tolist()


# =============================================================================
# PREENCHIMENTO DE FALHAS
# =============================================================================

def preencher_media(df):

    df2 = df.copy()

    for col in obter_colunas_numericas(df2):

        media = df2[col].mean()

        df2[col] = df2[col].fillna(media)

    return df2


def preencher_mediana(df):

    df2 = df.copy()

    for col in obter_colunas_numericas(df2):

        mediana = df2[col].median()

        df2[col] = df2[col].fillna(mediana)

    return df2


def preencher_moda(df):

    df2 = df.copy()

    for col in df2.columns:

        try:

            moda = df2[col].mode()[0]

            df2[col] = df2[col].fillna(moda)

        except:
            pass

    return df2


def preencher_interpolacao(df):

    df2 = df.copy()

    for col in obter_colunas_numericas(df2):

        try:

            df2[col] = df2[col].interpolate(
                method="linear",
                limit_direction="both"
            )

        except:
            pass

    return df2


def preencher_polinomial(df):

    df2 = df.copy()

    for col in obter_colunas_numericas(df2):

        try:

            df2[col] = df2[col].interpolate(
                method="polynomial",
                order=2
            )

        except:
            pass

    return df2


# =============================================================================
# REMOVER DUPLICADOS
# =============================================================================

def remover_duplicados(df):

    return df.drop_duplicates()


# =============================================================================
# PADRONIZAR NOMES DAS COLUNAS
# =============================================================================

def padronizar_colunas(df):

    df2 = df.copy()

    novas = []

    for col in df2.columns:

        nome = str(col)

        nome = nome.strip()

        nome = nome.replace(" ", "_")
        nome = nome.replace("-", "_")
        nome = nome.replace("/", "_")
        nome = nome.replace("(", "")
        nome = nome.replace(")", "")

        novas.append(nome)

    df2.columns = novas

    return df2


# =============================================================================
# DETECÇÃO DE OUTLIERS - IQR
# =============================================================================

def detectar_outliers_iqr(
    df,
    fator=1.5
):

    resultado = {}

    for col in obter_colunas_numericas(df):

        q1 = df[col].quantile(0.25)

        q3 = df[col].quantile(0.75)

        iqr = q3 - q1

        limite_inferior = (
            q1 - fator * iqr
        )

        limite_superior = (
            q3 + fator * iqr
        )

        mascara = (

            (df[col] < limite_inferior)

            |

            (df[col] > limite_superior)

        )

        resultado[col] = mascara

    return resultado


# =============================================================================
# DETECÇÃO DE OUTLIERS - Z SCORE
# =============================================================================

def detectar_outliers_zscore(
    df,
    limite=3
):

    resultado = {}

    for col in obter_colunas_numericas(df):

        try:

            z = np.abs(
                zscore(
                    df[col],
                    nan_policy="omit"
                )
            )

            resultado[col] = z > limite

        except:

            resultado[col] = np.zeros(
                len(df),
                dtype=bool
            )

    return resultado


# =============================================================================
# APLICAÇÃO DA AÇÃO NOS OUTLIERS
# =============================================================================

def aplicar_acao_outlier(
    df,
    mascaras,
    acao="remover"
):

    df2 = df.copy()

    mascara_total = np.zeros(
        len(df2),
        dtype=bool
    )

    for col in mascaras:

        mascara_total |= mascaras[col]

    if acao == "remover":

        df2 = df2[
            ~mascara_total
        ]

    elif acao == "media":

        for col in mascaras:

            media = df2[col].mean()

            df2.loc[
                mascaras[col],
                col
            ] = media

    elif acao == "mediana":

        for col in mascaras:

            mediana = df2[col].median()

            df2.loc[
                mascaras[col],
                col
            ] = mediana

    return df2


# =============================================================================
# ABA 2 - TRATAMENTO DE DADOS
# =============================================================================

with tab2:

    st.subheader(
        "🧹 Tratamento Inteligente dos Dados"
    )

    if "df_original" not in st.session_state:

        st.warning(
            "Importe uma planilha primeiro."
        )

    else:

        df_base = (
            st.session_state[
                "df_original"
            ].copy()
        )

        st.markdown(
            "### Valores Ausentes"
        )

        metodo_falhas = st.selectbox(

            "Método de preenchimento",

            [
                "Nenhum",
                "Média",
                "Mediana",
                "Moda",
                "Interpolação Linear",
                "Interpolação Polinomial"
            ]
        )

        st.markdown(
            "### Outliers"
        )

        metodo_outlier = st.selectbox(

            "Método de detecção",

            [
                "Nenhum",
                "IQR",
                "Z-Score"
            ]
        )

        acao_outlier = st.selectbox(

            "Ação para os outliers",

            [
                "remover",
                "media",
                "mediana"
            ]
        )

        remover_dup = st.checkbox(
            "Remover linhas duplicadas",
            value=True
        )

        padronizar = st.checkbox(
            "Padronizar nomes das colunas",
            value=True
        )

        st.markdown("---")

        if st.button(
            "🚀 Executar Tratamento",
            use_container_width=True
        ):

            df_tratado = df_base.copy()

            # ----------------------------------------------------
            # PADRONIZAÇÃO
            # ----------------------------------------------------

            if padronizar:

                df_tratado = (
                    padronizar_colunas(
                        df_tratado
                    )
                )

            # ----------------------------------------------------
            # DUPLICADOS
            # ----------------------------------------------------

            if remover_dup:

                antes = len(
                    df_tratado
                )

                df_tratado = (
                    remover_duplicados(
                        df_tratado
                    )
                )

                depois = len(
                    df_tratado
                )

                removidos = (
                    antes - depois
                )

                st.success(
                    f"{removidos} linhas duplicadas removidas."
                )

            # ----------------------------------------------------
            # FALHAS
            # ----------------------------------------------------

            if metodo_falhas == "Média":

                df_tratado = (
                    preencher_media(
                        df_tratado
                    )
                )

            elif metodo_falhas == "Mediana":

                df_tratado = (
                    preencher_mediana(
                        df_tratado
                    )
                )

            elif metodo_falhas == "Moda":

                df_tratado = (
                    preencher_moda(
                        df_tratado
                    )
                )

            elif metodo_falhas == "Interpolação Linear":

                df_tratado = (
                    preencher_interpolacao(
                        df_tratado
                    )
                )

            elif metodo_falhas == "Interpolação Polinomial":

                df_tratado = (
                    preencher_polinomial(
                        df_tratado
                    )
                )

            # ----------------------------------------------------
            # OUTLIERS
            # ----------------------------------------------------

            if metodo_outlier == "IQR":

                mascaras = (
                    detectar_outliers_iqr(
                        df_tratado
                    )
                )

                df_tratado = (
                    aplicar_acao_outlier(
                        df_tratado,
                        mascaras,
                        acao_outlier
                    )
                )

            elif metodo_outlier == "Z-Score":

                mascaras = (
                    detectar_outliers_zscore(
                        df_tratado
                    )
                )

                df_tratado = (
                    aplicar_acao_outlier(
                        df_tratado,
                        mascaras,
                        acao_outlier
                    )
                )

            st.session_state[
                "df_tratado"
            ] = df_tratado

            st.success(
                "Tratamento concluído com sucesso."
            )

        # ======================================================
        # RESULTADO
        # ======================================================

        if "df_tratado" in st.session_state:

            df_resultado = (
                st.session_state[
                    "df_tratado"
                ]
            )

            st.markdown("---")

            st.subheader(
                "📊 Resultado do Tratamento"
            )

            st.dataframe(
                df_resultado.head(100),
                use_container_width=True
            )

            col1, col2, col3 = st.columns(3)

            with col1:

                st.metric(
                    "Linhas",
                    len(df_resultado)
                )

            with col2:

                st.metric(
                    "Colunas",
                    len(
                        df_resultado.columns
                    )
                )

            with col3:

                st.metric(
                    "Valores Nulos",
                    int(
                        df_resultado
                        .isna()
                        .sum()
                        .sum()
                    )
                )

            st.success(
                "Base pronta para Consolidação."
            )
# =============================================================================
# FUNÇÕES DE CONSOLIDAÇÃO
# =============================================================================

def identificar_coluna_data(df):

    palavras = [
        "data",
        "date",
        "datetime",
        "tempo",
        "timestamp"
    ]

    for col in df.columns:

        nome = str(col).lower()

        if any(p in nome for p in palavras):
            return col

    return None


def consolidar_dataframe(df):

    df2 = df.copy()

    coluna_data = identificar_coluna_data(df2)

    if coluna_data is not None:

        try:

            df2 = df2.sort_values(
                by=coluna_data
            )

        except:
            pass

    df2 = df2.reset_index(
        drop=True
    )

    return df2


def gerar_resumo_consolidacao(df):

    resumo = {

        "Linhas":
            len(df),

        "Colunas":
            len(df.columns),

        "Valores Nulos":
            int(
                df.isna()
                .sum()
                .sum()
            ),

        "Colunas Numéricas":
            len(
                df.select_dtypes(
                    include=np.number
                ).columns
            )

    }

    return resumo
# =============================================================================
# ESTATÍSTICA DESCRITIVA
# =============================================================================

def classificar_cv(cv):

    if pd.isna(cv):
        return "-"

    if cv < 10:
        return "Baixo"

    elif cv < 20:
        return "Médio"

    elif cv < 30:
        return "Alto"

    else:
        return "Muito Alto"


def gerar_estatisticas(df):

    numericas = df.select_dtypes(
        include=np.number
    )

    resultados = []

    for col in numericas.columns:

        serie = numericas[col].dropna()

        if len(serie) == 0:
            continue

        media = serie.mean()

        desvio = serie.std()

        cv = (
            desvio / media * 100
            if media != 0
            else np.nan
        )

        try:
            moda = serie.mode().iloc[0]
        except:
            moda = np.nan

        resultados.append({

            "Variável": col,

            "N": len(serie),

            "Média": round(media, 4),

            "Mediana": round(
                serie.median(),
                4
            ),

            "Moda": round(
                moda,
                4
            ) if pd.notna(moda) else np.nan,

            "Mínimo": round(
                serie.min(),
                4
            ),

            "Máximo": round(
                serie.max(),
                4
            ),

            "Amplitude": round(
                serie.max() - serie.min(),
                4
            ),

            "Variância": round(
                serie.var(),
                4
            ),

            "Desvio Padrão": round(
                desvio,
                4
            ),

            "CV (%)": round(
                cv,
                2
            ),

            "Classificação CV":
                classificar_cv(cv),

            "Assimetria": round(
                serie.skew(),
                4
            ),

            "Curtose": round(
                serie.kurtosis(),
                4
            ),

            "P5": round(
                serie.quantile(0.05),
                4
            ),

            "P25": round(
                serie.quantile(0.25),
                4
            ),

            "P75": round(
                serie.quantile(0.75),
                4
            ),

            "P95": round(
                serie.quantile(0.95),
                4
            )

        })

    return pd.DataFrame(resultados)


# =============================================================================
# EVENTOS EXTREMOS
# =============================================================================

def detectar_extremos(df):

    extremos = []

    numericas = df.select_dtypes(
        include=np.number
    )

    for col in numericas.columns:

        try:

            maior = numericas[col].max()

            menor = numericas[col].min()

            extremos.append({

                "Variável": col,
                "Tipo": "Máximo",
                "Valor": maior

            })

            extremos.append({

                "Variável": col,
                "Tipo": "Mínimo",
                "Valor": menor

            })

        except:
            pass

    return pd.DataFrame(extremos)

# =============================================================================
# ABA 3 - CONSOLIDAÇÃO
# =============================================================================

with tab3:

    st.subheader(
        "📅 Consolidação dos Dados"
    )

    if "df_tratado" not in st.session_state:

        st.warning(
            "Primeiro execute o tratamento dos dados."
        )

    else:

        df_base = (
            st.session_state[
                "df_tratado"
            ].copy()
        )

        st.info(
            """
            Esta etapa organiza a base final,
            ordena datas e prepara os dados
            para análises estatísticas.
            """
        )

        if st.button(
            "🚀 Consolidar Dados",
            use_container_width=True
        ):

            df_consolidado = (
                consolidar_dataframe(
                    df_base
                )
            )

            st.session_state[
                "df_consolidado"
            ] = df_consolidado

            st.success(
                "Dados consolidados com sucesso."
            )

        if (
            "df_consolidado"
            in st.session_state
        ):

            df_consolidado = (
                st.session_state[
                    "df_consolidado"
                ]
            )

            resumo = (
                gerar_resumo_consolidacao(
                    df_consolidado
                )
            )

            st.markdown("---")

            st.subheader(
                "📊 Resumo da Consolidação"
            )

            col1, col2, col3, col4 = (
                st.columns(4)
            )

            with col1:

                st.metric(
                    "Linhas",
                    resumo["Linhas"]
                )

            with col2:

                st.metric(
                    "Colunas",
                    resumo["Colunas"]
                )

            with col3:

                st.metric(
                    "Nulos",
                    resumo[
                        "Valores Nulos"
                    ]
                )

            with col4:

                st.metric(
                    "Numéricas",
                    resumo[
                        "Colunas Numéricas"
                    ]
                )

            st.markdown("---")

            st.subheader(
                "👁 Pré-visualização"
            )

            st.dataframe(
                df_consolidado.head(200),
                use_container_width=True
            )

            st.success(
                "Base pronta para Estatística, Gráficos e IA."
            )
# =============================================================================
# ABA 4 - ESTATÍSTICA DESCRITIVA
# =============================================================================

with tab4:

    st.subheader(
        "📈 Estatística Descritiva"
    )

    if (
        "df_consolidado"
        not in st.session_state
    ):

        st.warning(
            "Primeiro consolide os dados."
        )

    else:

        df_base = st.session_state[
            "df_consolidado"
        ]

        estatisticas = gerar_estatisticas(
            df_base
        )

        extremos = detectar_extremos(
            df_base
        )

        st.markdown(
            "### Estatística Completa"
        )

        st.dataframe(
            estatisticas,
            use_container_width=True
        )

        st.markdown(
            "### Valores Extremos"
        )

        st.dataframe(
            extremos,
            use_container_width=True
        )

        col1, col2, col3 = st.columns(3)

        with col1:

            st.metric(
                "Variáveis",
                len(estatisticas)
            )

        with col2:

            st.metric(
                "Extremos",
                len(extremos)
            )

        with col3:

            st.metric(
                "Registros",
                len(df_base)
            )

        st.session_state[
            "estatisticas"
        ] = estatisticas

        csv_estatisticas = (
            estatisticas
            .to_csv(index=False)
            .encode("utf-8")
        )

        st.download_button(
            "📥 Baixar Estatísticas",
            csv_estatisticas,
            file_name="estatisticas_descritivas.csv",
            mime="text/csv",
            use_container_width=True
        )
# =============================================================================
# GRÁFICOS CIENTÍFICOS
# =============================================================================

def obter_numericas(df):

    return df.select_dtypes(
        include=np.number
    ).columns.tolist()


def obter_data(df):

    for col in df.columns:

        nome = str(col).lower()

        if any(
            termo in nome
            for termo in [
                "data",
                "date",
                "datetime",
                "tempo"
            ]
        ):
            return col

    return None


def grafico_temporal(
    df,
    variavel
):

    data_col = obter_data(df)

    if data_col is None:
        return None

    fig = px.line(
        df,
        x=data_col,
        y=variavel,
        markers=True,
        title=f"Série Temporal - {variavel}"
    )

    fig.update_layout(
        height=600
    )

    return fig


def grafico_histograma(
    df,
    variavel
):

    fig = px.histogram(
        df,
        x=variavel,
        nbins=30,
        marginal="box",
        title=f"Histograma - {variavel}"
    )

    fig.update_layout(
        height=600
    )

    return fig


def grafico_boxplot(
    df,
    variavel
):

    fig = px.box(
        df,
        y=variavel,
        points="outliers",
        title=f"Boxplot - {variavel}"
    )

    fig.update_layout(
        height=600
    )

    return fig


def grafico_violin(
    df,
    variavel
):

    fig = px.violin(
        df,
        y=variavel,
        box=True,
        points="all",
        title=f"Gráfico Violino - {variavel}"
    )

    fig.update_layout(
        height=600
    )

    return fig


def grafico_dispersao(
    df,
    x_var,
    y_var
):

    fig = px.scatter(
        df,
        x=x_var,
        y=y_var,
        trendline="ols",
        title=f"{x_var} × {y_var}"
    )

    fig.update_layout(
        height=650
    )

    return fig


def heatmap_correlacao(df):

    numericas = df.select_dtypes(
        include=np.number
    )

    if numericas.empty:
        return None

    corr = numericas.corr()

    fig = px.imshow(
        corr,
        text_auto=".2f",
        aspect="auto",
        color_continuous_scale="RdBu_r"
    )

    fig.update_layout(
        title="Matriz de Correlação",
        height=700
    )

    return fig


def grafico_tendencia(
    df,
    variavel
):

    data_col = obter_data(df)

    if data_col is None:
        return None

    fig = px.scatter(
        df,
        x=data_col,
        y=variavel,
        trendline="ols",
        title=f"Tendência Temporal - {variavel}"
    )

    fig.update_layout(
        height=650
    )

    return fig


# =============================================================================
# ABA 5 - GRÁFICOS CIENTÍFICOS
# =============================================================================

with tab5:

    st.subheader(
        "📊 Gráficos Científicos"
    )

    if (
        "df_consolidado"
        not in st.session_state
    ):

        st.warning(
            "Primeiro consolide os dados."
        )

    else:

        df_base = st.session_state[
            "df_consolidado"
        ]

        numericas = obter_numericas(
            df_base
        )

        if len(numericas) == 0:

            st.error(
                "Nenhuma variável numérica encontrada."
            )

        else:

            tipo = st.selectbox(
                "Tipo de gráfico",
                [
                    "Série Temporal",
                    "Histograma",
                    "Boxplot",
                    "Violino",
                    "Dispersão",
                    "Heatmap",
                    "Tendência"
                ]
            )

            if tipo == "Heatmap":

                fig = heatmap_correlacao(
                    df_base
                )

                if fig is not None:

                    st.plotly_chart(
                        fig,
                        use_container_width=True
                    )

            elif tipo == "Dispersão":

                col1, col2 = st.columns(2)

                with col1:

                    x_var = st.selectbox(
                        "Variável X",
                        numericas
                    )

                with col2:

                    y_var = st.selectbox(
                        "Variável Y",
                        numericas,
                        index=min(
                            1,
                            len(numericas) - 1
                        )
                    )

                fig = grafico_dispersao(
                    df_base,
                    x_var,
                    y_var
                )

                st.plotly_chart(
                    fig,
                    use_container_width=True
                )

            else:

                variavel = st.selectbox(
                    "Variável",
                    numericas
                )

                if tipo == "Série Temporal":

                    fig = grafico_temporal(
                        df_base,
                        variavel
                    )

                elif tipo == "Histograma":

                    fig = grafico_histograma(
                        df_base,
                        variavel
                    )

                elif tipo == "Boxplot":

                    fig = grafico_boxplot(
                        df_base,
                        variavel
                    )

                elif tipo == "Violino":

                    fig = grafico_violin(
                        df_base,
                        variavel
                    )

                elif tipo == "Tendência":

                    fig = grafico_tendencia(
                        df_base,
                        variavel
                    )

                else:

                    fig = None

                if fig is None:

                    st.warning(
                        "Nenhuma coluna de data foi encontrada."
                    )

                else:

                    st.plotly_chart(
                        fig,
                        use_container_width=True
                    )
import requests
import json
import os
GEMINI_API_KEY = "SUA_CHAVE_AQUI"

# =============================================================================
# IA - PREPARAÇÃO DE CONTEXTO
# =============================================================================

def gerar_contexto_automatico(df):

    numericas = df.select_dtypes(
        include=np.number
    )

    contexto = f"""
    Total de registros: {len(df)}
    Total de variáveis: {len(df.columns)}
    """

    for col in numericas.columns:

        try:

            contexto += f"""

Variável: {col}

Média: {numericas[col].mean():.2f}

Mediana: {numericas[col].median():.2f}

Mínimo: {numericas[col].min():.2f}

Máximo: {numericas[col].max():.2f}

Desvio padrão: {numericas[col].std():.2f}

Coeficiente de variação:
{(numericas[col].std()/numericas[col].mean()*100):.2f} %

"""

        except:
            pass

    return contexto


# =============================================================================
# PROMPTS ESPECIALIZADOS
# =============================================================================

PROMPTS = {

    "Relatório Científico":

    """
    Produza um relatório científico completo.

    Inclua:

    - Introdução dos dados
    - Interpretação estatística
    - Tendências observadas
    - Variabilidade
    - Eventos extremos
    - Conclusões técnicas

    Utilize linguagem científica.
    """,

    "Relatório Meteorológico":

    """
    Produza uma análise meteorológica detalhada.

    Avalie:

    - Temperatura
    - Precipitação
    - Umidade
    - Vento
    - Eventos extremos
    - Tendências climáticas
    """,

    "Relatório Agrícola":

    """
    Interprete os dados sob o ponto de vista agronômico.

    Avalie:

    - Disponibilidade hídrica
    - Riscos climáticos
    - Potencial produtivo
    - Conforto térmico
    - Impactos agrícolas
    """,

    "Resumo Executivo":

    """
    Produza um resumo executivo simples e objetivo.

    Explique os resultados em linguagem acessível.
    """
}


# =============================================================================
# CONSULTA GEMINI
# =============================================================================

def consultar_ia(
    prompt_usuario,
    contexto
):

    try:

        api_key = st.secrets[
            "GEMINI_API_KEY"
        ]

        url = (
            "https://generativelanguage.googleapis.com/"
            "v1beta/models/gemini-2.0-flash:generateContent"
        )

        payload = {

            "contents": [

                {

                    "parts": [

                        {

                            "text":
                            contexto
                            + "\n\n"
                            + prompt_usuario

                        }

                    ]

                }

            ]

        }

        resposta = requests.post(

            f"{url}?key={api_key}",

            headers={
                "Content-Type":
                "application/json"
            },

            json=payload,

            timeout=120

        )

        if resposta.status_code == 200:

            dados = resposta.json()

            return (
                dados["candidates"][0]
                ["content"]["parts"][0]
                ["text"]
            )

        return (
            f"Erro na API Gemini: "
            f"{resposta.status_code}"
        )

    except Exception as erro:

        return f"Erro: {erro}"


# =============================================================================
# ABA 6 - INTELIGÊNCIA ARTIFICIAL
# =============================================================================

with tab6:

    st.subheader(
        "🤖 Inteligência Artificial"
    )

    if (
        "df_consolidado"
        not in st.session_state
    ):

        st.warning(
            "Consolide os dados primeiro."
        )

    else:

        df_base = st.session_state[
            "df_consolidado"
        ]

        st.markdown(
            "### Relatórios Inteligentes"
        )

        tipo_relatorio = st.selectbox(

            "Tipo de Relatório",

            list(
                PROMPTS.keys()
            )
        )

        pergunta = st.text_area(

            "Pergunta adicional",

            placeholder=
            "Exemplo: Existe tendência de aumento da temperatura?"
        )

        if st.button(

            "🚀 Gerar Relatório",

            use_container_width=True

        ):

            with st.spinner(

                "Analisando dados..."

            ):

                contexto = (
                    gerar_contexto_automatico(
                        df_base
                    )
                )

                prompt = (

                    PROMPTS[
                        tipo_relatorio
                    ]

                    + "\n\n"

                    + pergunta

                )

                resposta = consultar_ia(

                    prompt,

                    contexto

                )

                st.session_state[
                    "relatorio_ia"
                ] = resposta

        if (

            "relatorio_ia"

            in st.session_state

        ):

            st.markdown("---")

            st.markdown(
                "### Relatório Gerado"
            )

            st.markdown(

                st.session_state[
                    "relatorio_ia"
                ]

            )

            st.download_button(

                "📥 Baixar Relatório TXT",

                st.session_state[
                    "relatorio_ia"
                ],

                file_name=
                "relatorio_ia.txt",

                mime="text/plain",

                use_container_width=True

            )
from io import BytesIO
import zipfile

from docx import Document
import pandas as pd
# =============================================================================
# EXPORTAÇÃO XLSX
# =============================================================================

def gerar_excel_completo():

    buffer = BytesIO()

    with pd.ExcelWriter(
        buffer,
        engine="openpyxl"
    ) as writer:

        if "df_tratado" in st.session_state:

            st.session_state[
                "df_tratado"
            ].to_excel(
                writer,
                sheet_name="Tratado",
                index=False
            )

        if "df_consolidado" in st.session_state:

            st.session_state[
                "df_consolidado"
            ].to_excel(
                writer,
                sheet_name="Consolidado",
                index=False
            )

        if "estatisticas" in st.session_state:

            st.session_state[
                "estatisticas"
            ].to_excel(
                writer,
                sheet_name="Estatisticas",
                index=False
            )

    buffer.seek(0)

    return buffer


# =============================================================================
# RELATÓRIO DOCX
# =============================================================================

def gerar_docx():

    doc = Document()

    doc.add_heading(
        "Relatório Técnico",
        level=0
    )

    doc.add_paragraph(
        "Documento gerado automaticamente pelo sistema."
    )

    if "relatorio_ia" in st.session_state:

        doc.add_heading(
            "Análise da Inteligência Artificial",
            level=1
        )

        doc.add_paragraph(
            st.session_state[
                "relatorio_ia"
            ]
        )

    arquivo = BytesIO()

    doc.save(arquivo)

    arquivo.seek(0)

    return arquivo


# =============================================================================
# RELATÓRIO HTML
# =============================================================================

def gerar_html():

    html = """

    <html>

    <head>

        <meta charset="utf-8">

        <title>
            Relatório Técnico
        </title>

    </head>

    <body>

        <h1>
            Relatório Técnico
        </h1>

    """

    if "relatorio_ia" in st.session_state:

        html += f"""

        <h2>
            Análise da IA
        </h2>

        <p>
            {st.session_state['relatorio_ia']}
        </p>

        """

    html += """

    </body>

    </html>

    """

    return html


# =============================================================================
# PACOTE ZIP
# =============================================================================

def gerar_zip():

    memoria = BytesIO()

    with zipfile.ZipFile(

        memoria,

        "w",

        zipfile.ZIP_DEFLATED

    ) as zipf:

        if "df_tratado" in st.session_state:

            zipf.writestr(

                "dados_tratados.csv",

                st.session_state[
                    "df_tratado"
                ].to_csv(
                    index=False
                )

            )

        if "df_consolidado" in st.session_state:

            zipf.writestr(

                "dados_consolidados.csv",

                st.session_state[
                    "df_consolidado"
                ].to_csv(
                    index=False
                )

            )

        if "estatisticas" in st.session_state:

            zipf.writestr(

                "estatisticas.csv",

                st.session_state[
                    "estatisticas"
                ].to_csv(
                    index=False
                )

            )

        if "relatorio_ia" in st.session_state:

            zipf.writestr(

                "relatorio_ia.txt",

                st.session_state[
                    "relatorio_ia"
                ]

            )

    memoria.seek(0)

    return memoria


# =============================================================================
# ABA 7 - EXPORTAÇÃO PROFISSIONAL
# =============================================================================

with tab7:

    st.subheader(
        "📥 Exportação Profissional"
    )

    st.info(
        "Exporte seus resultados em Excel, DOCX, HTML ou ZIP."
    )

    col1, col2 = st.columns(2)

    with col1:

        excel = gerar_excel_completo()

        st.download_button(

            "📊 Excel Completo",

            data=excel,

            file_name="analise_completa.xlsx",

            mime=(
                "application/"
                "vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),

            use_container_width=True

        )

        docx = gerar_docx()

        st.download_button(

            "📝 Relatório DOCX",

            data=docx,

            file_name="relatorio.docx",

            mime=(
                "application/"
                "vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),

            use_container_width=True

        )

    with col2:

        html = gerar_html()

        st.download_button(

            "🌐 Relatório HTML",

            data=html,

            file_name="relatorio.html",

            mime="text/html",

            use_container_width=True

        )

        zip_file = gerar_zip()

        st.download_button(

            "📦 Pacote Completo",

            data=zip_file,

            file_name="projeto_completo.zip",

            mime="application/zip",

            use_container_width=True

        )

    st.markdown("---")

    st.success(
        "Exportação pronta para Excel, Word, HTML e Backup ZIP."
    )
# =============================================================================
# INDICADORES AGROMETEOROLÓGICOS
# =============================================================================

def localizar_coluna(df, palavras):

    for col in df.columns:

        nome = str(col).lower()

        if any(
            p in nome
            for p in palavras
        ):
            return col

    return None


# =============================================================================
# GRAUS-DIA (GDD)
# =============================================================================

def calcular_gdd(
    df,
    temperatura,
    temperatura_base=10
):

    serie = pd.to_numeric(
        df[temperatura],
        errors="coerce"
    )

    gdd = (
        serie - temperatura_base
    )

    gdd[gdd < 0] = 0

    return gdd


# =============================================================================
# PRECIPITAÇÃO ACUMULADA
# =============================================================================

def calcular_precipitacao_acumulada(
    df,
    coluna
):

    chuva = pd.to_numeric(
        df[coluna],
        errors="coerce"
    )

    return chuva.cumsum()


# =============================================================================
# SOMA TÉRMICA
# =============================================================================

def calcular_soma_termica(gdd):

    return gdd.cumsum()


# =============================================================================
# BALANÇO HÍDRICO SIMPLIFICADO
# =============================================================================

def calcular_balanco_hidrico(
    precipitacao,
    eto
):

    return precipitacao - eto


# =============================================================================
# ETo SIMPLIFICADA
# Hargreaves simplificada
# =============================================================================

def calcular_eto_hargreaves(
    tmin,
    tmax,
    tmed
):

    eto = (

        0.0023

        *

        (tmed + 17.8)

        *

        np.sqrt(tmax - tmin)

    )

    return eto


# =============================================================================
# ÍNDICE DE CONFORTO TÉRMICO
# =============================================================================

def indice_conforto_termico(
    temperatura,
    umidade
):

    return (

        temperatura

        -

        (
            0.55
            -
            0.0055 * umidade
        )

        *

        (
            temperatura - 14.5
        )

    )
# =============================================================================
# ABA 8 - INDICADORES AGROMETEOROLÓGICOS
# =============================================================================

with tab8:

    st.subheader(
        "🌱 Indicadores Agrometeorológicos"
    )

    if (
        "df_consolidado"
        not in st.session_state
    ):

        st.warning(
            "Consolide os dados primeiro."
        )

    else:

        df = st.session_state[
            "df_consolidado"
        ]

        st.markdown(
            """
            Cálculo automático de indicadores
            agrícolas e meteorológicos.
            """
        )

        temperatura = localizar_coluna(

            df,

            [
                "temp",
                "temperatura",
                "tmed"
            ]

        )

        precipitacao = localizar_coluna(

            df,

            [
                "chuva",
                "precip",
                "precipitacao"
            ]

        )

        umidade = localizar_coluna(

            df,

            [
                "umidade",
                "ur"
            ]

        )

        tmin = localizar_coluna(
            df,
            ["tmin"]
        )

        tmax = localizar_coluna(
            df,
            ["tmax"]
        )

        if temperatura is not None:

            temperatura_base = st.number_input(

                "Temperatura Base (°C)",

                value=10.0

            )

            gdd = calcular_gdd(

                df,

                temperatura,

                temperatura_base

            )

            soma_termica = (
                calcular_soma_termica(
                    gdd
                )
            )

            st.metric(

                "Graus-dia Acumulados",

                round(
                    soma_termica.iloc[-1],
                    2
                )

            )

            df["GDD"] = gdd

            df["Soma_Termica"] = (
                soma_termica
            )

        if precipitacao is not None:

            chuva_acumulada = (

                calcular_precipitacao_acumulada(

                    df,

                    precipitacao

                )

            )

            df[
                "Chuva_Acumulada"
            ] = chuva_acumulada

            st.metric(

                "Precipitação Acumulada (mm)",

                round(
                    chuva_acumulada.iloc[-1],
                    2
                )

            )

        if (

            tmin is not None

            and

            tmax is not None

            and

            temperatura is not None

        ):

            eto = calcular_eto_hargreaves(

                df[tmin],

                df[tmax],

                df[temperatura]

            )

            df["ETo"] = eto

            st.metric(

                "ETo Média",

                round(
                    eto.mean(),
                    2
                )

            )

        if (

            temperatura is not None

            and

            umidade is not None

        ):

            conforto = (

                indice_conforto_termico(

                    df[temperatura],

                    df[umidade]

                )

            )

            df[
                "Indice_Conforto"
            ] = conforto

            st.metric(

                "Conforto Médio",

                round(
                    conforto.mean(),
                    2
                )

            )

        st.markdown("---")

        st.subheader(
            "📋 Indicadores Gerados"
        )

        st.dataframe(
            df.head(100),
            use_container_width=True
        )

        st.session_state[
            "df_indicadores"
        ] = df
# =============================================================================
# CONTROLE DE QUALIDADE METEOROLÓGICA
# =============================================================================

def localizar_coluna_por_nome(df, palavras):

    for col in df.columns:

        nome = str(col).lower()

        if any(
            p in nome
            for p in palavras
        ):
            return col

    return None


# =============================================================================
# VERIFICAÇÃO DE LIMITES FÍSICOS
# =============================================================================

def verificar_limites_fisicos(df):

    problemas = []

    temp = localizar_coluna_por_nome(
        df,
        ["temp", "temperatura"]
    )

    ur = localizar_coluna_por_nome(
        df,
        ["umidade", "ur"]
    )

    chuva = localizar_coluna_por_nome(
        df,
        ["chuva", "precip"]
    )

    vento = localizar_coluna_por_nome(
        df,
        ["vento", "wind"]
    )

    if temp is not None:

        erros = (
            (df[temp] < -20)
            |
            (df[temp] > 60)
        )

        problemas.append({

            "Variável": temp,
            "Registros Suspeitos":
            int(erros.sum())
        })

    if ur is not None:

        erros = (
            (df[ur] < 0)
            |
            (df[ur] > 100)
        )

        problemas.append({

            "Variável": ur,
            "Registros Suspeitos":
            int(erros.sum())
        })

    if chuva is not None:

        erros = (
            df[chuva] < 0
        )

        problemas.append({

            "Variável": chuva,
            "Registros Suspeitos":
            int(erros.sum())
        })

    if vento is not None:

        erros = (
            df[vento] < 0
        )

        problemas.append({

            "Variável": vento,
            "Registros Suspeitos":
            int(erros.sum())
        })

    return pd.DataFrame(
        problemas
    )


# =============================================================================
# RELATÓRIO DE FALHAS
# =============================================================================

def gerar_relatorio_falhas(df):

    relatorio = pd.DataFrame({

        "Coluna":
            df.columns,

        "Valores Nulos":
            df.isna().sum(),

        "Percentual (%)":
            (
                df.isna().sum()
                /
                len(df)
                *
                100
            ).round(2)

    })

    return relatorio


# =============================================================================
# SCORE DE QUALIDADE
# =============================================================================

def calcular_score_qualidade(df):

    total = (
        len(df)
        *
        len(df.columns)
    )

    faltantes = (
        df.isna()
        .sum()
        .sum()
    )

    score = (

        (
            total
            -
            faltantes
        )

        /

        total

    ) * 100

    return round(
        score,
        2
    )
# =============================================================================
# ABA 9 - QUALIDADE DOS DADOS
# =============================================================================

with tab9:

    st.subheader(
        "🔍 Controle de Qualidade Meteorológica"
    )

    if (
        "df_consolidado"
        not in st.session_state
    ):

        st.warning(
            "Consolide os dados primeiro."
        )

    else:

        df = st.session_state[
            "df_consolidado"
        ]

        score = (
            calcular_score_qualidade(
                df
            )
        )

        st.metric(
            "Score de Qualidade (%)",
            score
        )

        st.markdown("---")

        st.subheader(
            "Valores Ausentes"
        )

        falhas = (
            gerar_relatorio_falhas(
                df
            )
        )

        st.dataframe(
            falhas,
            use_container_width=True
        )

        st.markdown("---")

        st.subheader(
            "Limites Físicos"
        )

        limites = (
            verificar_limites_fisicos(
                df
            )
        )

        st.dataframe(
            limites,
            use_container_width=True
        )

        if score >= 95:

            st.success(
                "Base considerada excelente."
            )

        elif score >= 85:

            st.info(
                "Base considerada boa."
            )

        elif score >= 70:

            st.warning(
                "Base necessita revisão."
            )

        else:

            st.error(
                "Base apresenta baixa qualidade."
            )
# =============================================================================
# RODAPÉ INSTITUCIONAL
# =============================================================================

st.markdown("---")

st.markdown(
    """
    <div style='text-align:center;
                font-size:13px;
                color:gray;
                padding-top:10px;
                padding-bottom:20px;'>

    <b>AgroClimate AI</b><br>

    Sistema experimental desenvolvido para fins acadêmicos,
    educacionais, científicos e de apoio à análise de dados
    meteorológicos e agrometeorológicos.<br><br>

    Este aplicativo não substitui análises técnicas oficiais,
    laudos periciais, pareceres especializados ou sistemas
    operacionais de instituições governamentais.<br><br>

    Parte dos métodos, conceitos e indicadores utilizados baseia-se
    em literatura científica, manuais técnicos e referências
    amplamente adotadas na área de Meteorologia, Climatologia,
    Agrometeorologia e Ciências Agrárias, incluindo publicações do
    Instituto Nacional de Meteorologia (INMET), da Organização
    Meteorológica Mundial (OMM/WMO), da FAO e de trabalhos
    científicos especializados.<br><br>

    Os resultados gerados devem ser interpretados por profissionais
    qualificados e utilizados apenas como ferramenta complementar
    de apoio à tomada de decisão.<br><br>

    © 2026 AgroClimate AI — Versão Acadêmica Experimental

    </div>
    """,
    unsafe_allow_html=True
)
